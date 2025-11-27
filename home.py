import streamlit as st
from docx import Document
from io import BytesIO
import time

def read_docx_form(uploaded_file):
    # Load DOCX from uploaded file bytes
    doc = Document(BytesIO(uploaded_file.read()))
    results = []
    form_data = {}

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            results.append(para.text.strip())

    # Extract table contents
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    results.append(text)

    # Convert "Key: Value" lines to dict
    for line in results:
        if ":" in line:
            key, value = line.split(":", 1)
            form_data[key.strip()] = value.strip()

    return form_data


def chat_stream(prompt):
    response = f'You said, "{prompt}" ...interesting.'
    for char in response:
        yield char
        time.sleep(0.02)


def save_feedback(index):
    st.session_state.history[index]["feedback"] = st.session_state[f"feedback_{index}"]


st.title("Document Extractor")

# Initialize chat history
if "history" not in st.session_state:
    st.session_state.history = []


# Display chat history
for i, message in enumerate(st.session_state.history):
    with st.chat_message(message["role"]):
        st.write(message["content"])
        
        if message["role"] == "assistant":
            feedback = message.get("feedback", None)
            st.session_state[f"feedback_{i}"] = feedback
            
            st.feedback(
                "thumbs",
                key=f"feedback_{i}",
                disabled=feedback is not None,
                on_change=save_feedback,
                args=[i],
            )


# NEW: Upload DOCX instead of chat input
uploaded_doc = st.file_uploader("Upload DOCX file", type=["docx"])

if uploaded_doc:
    with st.chat_message("user"):
        st.write(f"Uploaded file: **{uploaded_doc.name}**")

    st.session_state.history.append(
        {"role": "user", "content": f"Uploaded {uploaded_doc.name}"}
    )

    # Extract content
    extracted = read_docx_form(uploaded_doc)

    with st.chat_message("assistant"):
        st.write("📄 Extracted Form Data:")
        st.table(extracted)

        st.feedback(
            "thumbs",
            key=f"feedback_{len(st.session_state.history)}",
            on_change=save_feedback,
            args=[len(st.session_state.history)],
        )

    st.session_state.history.append(
        {"role": "assistant", "content": str(extracted)}
    )
