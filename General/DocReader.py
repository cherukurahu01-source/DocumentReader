from docx import Document

def read_docx_form(file_path):
    doc = Document(file_path)   # <-- Correct!
    results = []
    form_data = {}

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            results.append(para.text.strip())

    # Extract tables (cell by cell)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    results.append(text)
                    
    for line in results:
        if ":" in line:
            key, value = line.split(":",1)
            form_data[key.strip()] = value.strip()
            

    return form_data

file_path = r"C:\Users\cherur\Downloads\RIA_GDL_STU_Bug_Form.docx"
content = read_docx_form(file_path)

form_data = {}
    
for line in content:
    if ":" in line:
        key, value = line.split(":", 1)
        form_data[key.strip()] = value.strip()
        
print(form_data)